import re
import uuid
import json
import os
from docx import Document
from typing import Dict, Any, List
from langchain_community.graphs.graph_document import GraphDocument, Node, Relationship
from langchain_core.documents import Document

def gen_id(prefix):
    return f"{prefix}_{uuid.uuid4().hex[:8]}"

def parse_docx_to_json(docx_path: str) -> dict:
    doc = Document(docx_path)

    data = {
        "Documents": [],
        "Chapters": [],
        "Articles": [],
        "Sections": [],
        "Subsections": [],
        "Procedures": [],
        "Steps": [],
        "Relationships": []
    }

    # --- Metadata Document ---
    doc_id = gen_id("doc")
    document_meta = {
        "id": doc_id,
        "Name": os.path.basename(docx_path),
        "Number_of_Pages": None,
        "Number_of_Chapter": 0
    }
    data["Documents"].append(document_meta)

    current_chapter_id = None
    current_section_id = None
    current_procedure_id = None
    current_article_id = None

    # --- Duyệt paragraph ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # === Detect CHAPTER ===
        if re.match(r"^CHƯƠNG\s+\w+", text.upper()):
            chap_id = gen_id("chapter")
            number_match = re.search(r"CHƯƠNG\s+([IVXLC0-9]+)", text.upper())
            chap_number = number_match.group(1) if number_match else None
            chapter_name = text

            chapter = {"id": chap_id, "Name": chapter_name, "Number": chap_number}
            data["Chapters"].append(chapter)
            data["Relationships"].append({"from": doc_id, "to": chap_id, "type": "HAS_CHAPTER"})
            document_meta["Number_of_Chapter"] += 1
            current_chapter_id = chap_id
            current_section_id = None
            current_procedure_id = None
            continue

        # === Detect SECTION ===
        if re.match(r"^\d+(\.| )", text) and not re.match(r"^\d+\.\d+", text):
            if current_article_id:
                continue
            sec_id = gen_id("section")
            sec_number = re.match(r"^(\d+)", text).group(1)
            section = {"id": sec_id, "Name": text, "Number": int(sec_number)}
            data["Sections"].append(section)
            if current_chapter_id:
                data["Relationships"].append({"from": current_chapter_id, "to": sec_id, "type": "HAS_SECTION"})
            current_section_id = sec_id
            current_procedure_id = None
            continue

        # === Detect SUBSECTION ===
        if re.match(r"^[a-z]\)", text) or re.match(r"^[a-z]\.", text):
            if not current_section_id:
                continue
            sub_id = gen_id("subsec")
            sub_letter = re.match(r"^([a-z])", text).group(1)
            subsection = {"id": sub_id, "Name": text, "Letter": sub_letter}
            data["Subsections"].append(subsection)
            data["Relationships"].append({"from": current_section_id, "to": sub_id, "type": "CONTAINS"})
            continue

        # === Detect PROCEDURE ===
        if re.match(r"^QUY TRÌNH\s+\d+", text.upper()):
            proc_id = gen_id("proc")
            number_match = re.search(r"QUY TRÌNH\s+(\d+)", text.upper())
            proc_number = number_match.group(1) if number_match else None
            procedure = {"id": proc_id, "Name": text, "Number": int(proc_number) if proc_number else None}
            data["Procedures"].append(procedure)
            if current_chapter_id:
                data["Relationships"].append({"from": current_chapter_id, "to": proc_id, "type": "HAS_PROCEDURE"})
            current_procedure_id = proc_id
            continue

        # === Detect STEP ===
        if re.match(r"^\d+\.\d+", text):
            step_id = gen_id("step")
            step_number = re.match(r"^(\d+(\.\d+)+)", text).group(1)
            step = {"id": step_id, "Name": text, "Number": step_number}
            data["Steps"].append(step)
            if current_procedure_id:
                data["Relationships"].append({"from": current_procedure_id, "to": step_id, "type": "HAS_STEP"})
            continue

        # === Detect ARTICLE ===
        if re.match(r"^\s*Điều\s+\d+(\.|)", text, re.IGNORECASE):
            art_id = gen_id("article")
            number_match = re.search(r"Điều\s+(\d+)", text, re.IGNORECASE)
            art_number = int(number_match.group(1)) if number_match else None
            article = {"id": art_id, "Name": text, "Number": art_number}
            data["Articles"].append(article)
            if current_chapter_id:
                data["Relationships"].append({"from": current_chapter_id, "to": art_id, "type": "HAS_ARTICLE"})
            current_article_id = art_id
            current_section_id = None
            current_procedure_id = None
            continue

    return data

def build_graph_documents(payload: Dict[str, Any]) -> List[GraphDocument]:
    """Convert parsed JSON payload thành GraphDocuments để đẩy lên Neo4j"""
    nodes: Dict[str, Node] = {}
    rels: List[Relationship] = []

    # Helper tạo node
    def add_nodes(label: str, items: list, props_map: Dict[str,str]):
        for it in items:
            nid = it["id"]
            props = {pname: it.get(jname) for pname, jname in props_map.items()}
            nodes[nid] = Node(type=label, id=nid, properties=props)

    # Tạo các loại node
    add_nodes("Document", payload.get("Documents", []), {
        "name": "Name",
        "number_of_pages": "Number_of_Pages",
        "number_of_chapter": "Number_of_Chapter"
    })
    add_nodes("Chapter", payload.get("Chapters", []), {
        "name": "Name",
        "number": "Number"
    })
    add_nodes("Article", payload.get("Articles", []), {
        "name": "Name",
        "number": "Number"
    })
    add_nodes("Section", payload.get("Sections", []), {
        "name": "Name",
        "number": "Number"
    })
    add_nodes("Subsection", payload.get("Subsections", []), {
        "name": "Name",
        "letter": "Letter"
    })
    add_nodes("Procedure", payload.get("Procedures", []), {
        "name": "Name",
        "number": "Number"
    })
    add_nodes("Step", payload.get("Steps", []), {
        "name": "Name",
        "number": "Number"
    })

    # Tạo quan hệ
    for rel in payload.get("Relationships", []):
        src = nodes.get(rel["from"])
        tgt = nodes.get(rel["to"])
        if src and tgt:
            rels.append(Relationship(source=src, target=tgt, type=rel["type"], properties={}))

    # GraphDocument để import vào Neo4j
    src_doc = Document(page_content="Parsed DOCX", metadata={})
    return [GraphDocument(nodes=list(nodes.values()), relationships=rels, source=src_doc)]



